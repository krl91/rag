"""
Extracteur de fichiers bureautiques — Word (.docx), Excel (.xlsx), PDF.

- Word  : paragraphes, titres (→ Markdown headings), tableaux (→ Markdown).
- Excel : une entrée par feuille, données converties en tableau Markdown.
- PDF   : texte par page, numéro de page conservé dans les métadonnées.
"""

import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Union

from kb_smart_metering.normalize.models import RawDocument, normalize_file_source_id

logger = logging.getLogger(__name__)


class WordExtractor:
    """Extracteur de documents Word (.docx) via python-docx."""

    def __init__(self, file_path: Union[str, Path]) -> None:
        """
        Args:
            file_path: Chemin vers le fichier .docx à extraire.
        """
        self._file_path = Path(file_path)

    def extract(self) -> list[RawDocument]:
        """Extrait paragraphes, titres et tableaux du fichier Word."""
        import docx  # python-docx

        try:
            doc = docx.Document(self._file_path)
        except Exception as exc:
            logger.error("Impossible d'ouvrir %s : %s", self._file_path, exc)
            return []

        parts: list[str] = []

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                level_str = style_name.replace("Heading", "").strip()
                level = int(level_str) if level_str.isdigit() else 1
                parts.append(f"{'#' * level} {para.text}")
            else:
                parts.append(para.text)

        for table in doc.tables:
            table_lines: list[str] = []
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
                table_lines.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    table_lines.append("| " + " | ".join("---" for _ in cells) + " |")
            if table_lines:
                parts.append("\n".join(table_lines))

        contenu = "\n\n".join(parts)
        core_props = doc.core_properties
        stat = self._file_path.stat()

        logger.info(
            "Word : %d paragraphes, %d tableaux — %s",
            len(doc.paragraphs),
            len(doc.tables),
            self._file_path.name,
        )
        return [
            RawDocument(
                id_source=normalize_file_source_id(self._file_path),
                source_type="docx",
                titre=self._file_path.stem,
                contenu_texte=contenu,
                auteur=core_props.author or None,
                date_creation=core_props.created,
                date_modification=core_props.modified,
                url_ou_chemin=self._file_path.resolve().as_uri(),
                metadonnees={
                    "nb_paragraphes": len(doc.paragraphs),
                    "nb_tableaux": len(doc.tables),
                    "taille_octets": stat.st_size,
                },
            )
        ]


class ExcelExtractor:
    """Extracteur de classeurs Excel (.xlsx) via openpyxl."""

    def __init__(self, file_path: Union[str, Path]) -> None:
        """
        Args:
            file_path: Chemin vers le fichier .xlsx à extraire.
        """
        self._file_path = Path(file_path)

    def extract(self) -> list[RawDocument]:
        """Extrait une entrée par feuille du classeur Excel."""
        import openpyxl

        try:
            wb = openpyxl.load_workbook(self._file_path, read_only=True, data_only=True)
        except Exception as exc:
            logger.error("Impossible d'ouvrir %s : %s", self._file_path, exc)
            return []

        documents: list[RawDocument] = []
        stat = self._file_path.stat()

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))

            if not rows:
                continue

            markdown_rows: list[str] = []
            for i, row in enumerate(rows):
                cells = [str(cell) if cell is not None else "" for cell in row]
                markdown_rows.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    markdown_rows.append("| " + " | ".join("---" for _ in cells) + " |")

            contenu = f"## Feuille : {sheet_name}\n\n" + "\n".join(markdown_rows)
            nb_cols = max((len(r) for r in rows), default=0)

            documents.append(
                RawDocument(
                    id_source=f"{normalize_file_source_id(self._file_path)}#{sheet_name}",
                    source_type="xlsx",
                    titre=f"{self._file_path.stem} — {sheet_name}",
                    contenu_texte=contenu,
                    url_ou_chemin=self._file_path.resolve().as_uri(),
                    metadonnees={
                        "sheet_name": sheet_name,
                        "nb_lignes": len(rows),
                        "nb_colonnes": nb_cols,
                        "taille_octets": stat.st_size,
                    },
                )
            )

        wb.close()
        logger.info(
            "Excel : %d feuille(s) extraite(s) — %s", len(documents), self._file_path.name
        )
        return documents


class PdfExtractor:
    """Extracteur de documents PDF via PyMuPDF (fitz)."""

    def __init__(self, file_path: Union[str, Path]) -> None:
        """
        Args:
            file_path: Chemin vers le fichier .pdf à extraire.
        """
        self._file_path = Path(file_path)

    def extract(self) -> list[RawDocument]:
        """Extrait le texte de chaque page du PDF avec son numéro de page."""
        import fitz  # PyMuPDF

        try:
            pdf_doc = fitz.open(self._file_path)
        except Exception as exc:
            logger.error("Impossible d'ouvrir %s : %s", self._file_path, exc)
            return []

        documents: list[RawDocument] = []
        stat = self._file_path.stat()
        titre_base = self._file_path.stem
        metadata = pdf_doc.metadata or {}
        auteur = metadata.get("author") or None

        for page_num in range(pdf_doc.page_count):
            page = pdf_doc[page_num]
            texte = page.get_text()

            if not texte.strip():
                continue

            documents.append(
                RawDocument(
                    id_source=f"{normalize_file_source_id(self._file_path)}#page{page_num + 1}",
                    source_type="pdf",
                    titre=f"{titre_base} — page {page_num + 1}",
                    contenu_texte=f"[Page {page_num + 1}]\n{texte}",
                    auteur=auteur,
                    date_creation=_parse_pdf_date(metadata.get("creationDate")),
                    date_modification=_parse_pdf_date(metadata.get("modDate")),
                    url_ou_chemin=self._file_path.resolve().as_uri(),
                    metadonnees={
                        "page": page_num + 1,
                        "total_pages": pdf_doc.page_count,
                        "taille_octets": stat.st_size,
                        "pdf_title": metadata.get("title") or None,
                        "pdf_subject": metadata.get("subject") or None,
                    },
                )
            )

        pdf_doc.close()
        logger.info(
            "PDF : %d page(s) extraite(s) — %s", len(documents), self._file_path.name
        )
        return documents


def _parse_pdf_date(value: str | None) -> datetime | None:
    """Parse une date PDF au format D:YYYYMMDDHHmmSS[±HH'MM']."""
    if not value:
        return None
    clean = value[2:] if value.startswith("D:") else value
    digits = "".join(c for c in clean[:14] if c.isdigit())
    if len(digits) < 8:
        return None
    try:
        if len(digits) >= 14:
            return datetime.strptime(digits[:14], "%Y%m%d%H%M%S").replace(tzinfo=timezone.utc)
        return datetime.strptime(digits[:8], "%Y%m%d").replace(tzinfo=timezone.utc)
    except ValueError:
        return None
