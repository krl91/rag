"""
Tests unitaires des extracteurs Phase 1.

- API distantes (Jira, Confluence) : mockées via unittest.mock.
- Dépôts Git : créés dans des répertoires temporaires.
- Fichiers bureautiques : générés en mémoire avec les bibliothèques réelles.
- Transcriptions : lues depuis tests/fixtures/ ou créées dans tmp_path.
"""

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from kb_smart_metering.extractors.confluence import ConfluenceExtractor
from kb_smart_metering.extractors.git import GitExtractor
from kb_smart_metering.extractors.jira import JiraExtractor
from kb_smart_metering.extractors.meetings import MeetingExtractor
from kb_smart_metering.extractors.office import ExcelExtractor, PdfExtractor, WordExtractor
from kb_smart_metering.normalize.models import RawDocument

FIXTURES_DIR = Path(__file__).parent / "fixtures"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def assert_raw_document(doc: RawDocument, expected_source_type: str) -> None:
    assert isinstance(doc, RawDocument), f"Attendu RawDocument, reçu {type(doc)}"
    assert doc.source_type == expected_source_type
    assert doc.id_source
    assert doc.titre
    assert isinstance(doc.contenu_texte, str)
    assert isinstance(doc.metadonnees, dict)
    assert isinstance(doc.pieces_jointes, list)


# ---------------------------------------------------------------------------
# JiraExtractor
# ---------------------------------------------------------------------------

_JIRA_ISSUE = {
    "key": "SMART-42",
    "fields": {
        "summary": "Intégration MDM v3",
        "description": "Description du ticket de test.",
        "status": {"name": "In Progress"},
        "assignee": {"displayName": "Jean Dupont"},
        "reporter": {"displayName": "Marie Curie"},
        "created": "2024-01-10T09:00:00.000+0100",
        "updated": "2024-01-15T14:30:00.000+0100",
        "comment": {
            "comments": [
                {"author": {"displayName": "Jean Dupont"}, "body": "Premier commentaire."},
            ]
        },
        "issuelinks": [
            {"type": {"name": "blocks"}, "outwardIssue": {"key": "SMART-43"}},
        ],
    },
    "changelog": {
        "histories": [
            {
                "created": "2024-01-12T10:00:00.000+0100",
                "items": [
                    {"field": "status", "fromString": "Open", "toString": "In Progress"}
                ],
            }
        ]
    },
}


@patch("kb_smart_metering.extractors.jira.Jira")
def test_jira_extract_retourne_raw_document(mock_cls: MagicMock) -> None:
    mock_jira = MagicMock()
    mock_jira.url = "https://jira.example.com"
    mock_jira.jql.return_value = {"issues": [_JIRA_ISSUE], "total": 1}
    mock_cls.return_value = mock_jira

    docs = JiraExtractor(url="https://jira.example.com", token="s", jql="project=SMART").extract()

    assert len(docs) == 1
    doc = docs[0]
    assert_raw_document(doc, "jira")
    assert doc.id_source == "SMART-42"
    assert "Intégration MDM v3" in doc.titre
    assert "Description du ticket" in doc.contenu_texte
    assert "Premier commentaire" in doc.contenu_texte
    assert "Open → In Progress" in doc.contenu_texte
    assert doc.metadonnees["nb_comments"] == 1
    assert doc.metadonnees["nb_status_changes"] == 1
    assert doc.metadonnees["nb_links"] == 1
    assert doc.auteur == "Marie Curie"


@patch("kb_smart_metering.extractors.jira.Jira")
def test_jira_pagination_stop_sur_liste_vide(mock_cls: MagicMock) -> None:
    mock_jira = MagicMock()
    mock_jira.url = "https://jira.example.com"
    mock_jira.jql.side_effect = [
        {"issues": [_JIRA_ISSUE], "total": 1},
        {"issues": [], "total": 1},
    ]
    mock_cls.return_value = mock_jira

    docs = JiraExtractor(url="https://jira.example.com", token="s").extract()
    assert len(docs) == 1
    assert mock_jira.jql.call_count == 1  # stop dès que total atteint


@patch("kb_smart_metering.extractors.jira.Jira")
def test_jira_erreur_api_retourne_liste_vide(mock_cls: MagicMock) -> None:
    mock_jira = MagicMock()
    mock_jira.url = "https://jira.example.com"
    mock_jira.jql.side_effect = Exception("Connexion refusée")
    mock_cls.return_value = mock_jira

    docs = JiraExtractor(url="https://jira.example.com", token="s").extract()
    assert docs == []


# ---------------------------------------------------------------------------
# ConfluenceExtractor
# ---------------------------------------------------------------------------

_CONFLUENCE_PAGE = {
    "id": "123456",
    "title": "Architecture Smart Metering",
    "body": {
        "storage": {
            "value": "<h1>Architecture</h1><p>Description de l'architecture système.</p>"
        }
    },
    "version": {"number": 3, "when": "2024-02-01T10:00:00.000Z"},
    "history": {
        "createdBy": {"displayName": "Alice Martin"},
        "createdDate": "2024-01-01T09:00:00.000Z",
    },
    "space": {"key": "ARCH"},
    "metadata": {
        "labels": {"results": [{"name": "architecture"}, {"name": "smart-metering"}]}
    },
}


@patch("kb_smart_metering.extractors.confluence.Confluence")
def test_confluence_extract_retourne_raw_document(mock_cls: MagicMock) -> None:
    mock_conf = MagicMock()
    mock_conf.url = "https://confluence.example.com"
    mock_conf.get_all_pages_from_space.return_value = [_CONFLUENCE_PAGE]
    mock_conf.get_page_comments.return_value = {"results": []}
    mock_cls.return_value = mock_conf

    docs = ConfluenceExtractor(
        url="https://confluence.example.com", token="s", space_key="ARCH"
    ).extract()

    assert len(docs) == 1
    doc = docs[0]
    assert_raw_document(doc, "confluence")
    assert doc.id_source == "123456"
    assert "Architecture Smart Metering" in doc.titre
    assert "Architecture" in doc.contenu_texte
    assert "<h1>" not in doc.contenu_texte
    assert doc.auteur == "Alice Martin"
    assert doc.metadonnees["space_key"] == "ARCH"
    assert "architecture" in doc.metadonnees["labels"]


@patch("kb_smart_metering.extractors.confluence.Confluence")
def test_confluence_storage_format_converti_en_texte(mock_cls: MagicMock) -> None:
    """Le HTML du storage format est entièrement converti en texte brut."""
    page = {**_CONFLUENCE_PAGE, "body": {"storage": {"value": "<h2>Titre</h2><p>Corps <b>important</b>.</p>"}}}
    mock_conf = MagicMock()
    mock_conf.url = "https://confluence.example.com"
    mock_conf.get_all_pages_from_space.return_value = [page]
    mock_conf.get_page_comments.return_value = {"results": []}
    mock_cls.return_value = mock_conf

    docs = ConfluenceExtractor(url="https://confluence.example.com", token="s", space_key="X").extract()
    assert len(docs) == 1
    assert "<h2>" not in docs[0].contenu_texte
    assert "Titre" in docs[0].contenu_texte
    assert "important" in docs[0].contenu_texte


@patch("kb_smart_metering.extractors.confluence.Confluence")
def test_confluence_erreur_api_retourne_liste_vide(mock_cls: MagicMock) -> None:
    mock_conf = MagicMock()
    mock_conf.url = "https://confluence.example.com"
    mock_conf.get_all_pages_from_space.side_effect = Exception("Timeout")
    mock_cls.return_value = mock_conf

    docs = ConfluenceExtractor(url="https://confluence.example.com", token="s", space_key="X").extract()
    assert docs == []


# ---------------------------------------------------------------------------
# GitExtractor
# ---------------------------------------------------------------------------


def _init_git_repo(path: Path) -> "git.Repo":  # type: ignore[name-defined]
    import git

    repo = git.Repo.init(path)
    with repo.config_writer() as cw:
        cw.set_value("user", "name", "Test User")
        cw.set_value("user", "email", "test@example.com")
    return repo


def test_git_extract_commits(tmp_path: Path) -> None:
    repo = _init_git_repo(tmp_path)

    f = tmp_path / "README.md"
    f.write_text("# Projet\nContenu initial.")
    repo.index.add(["README.md"])
    repo.index.commit("feat: initialisation du projet")

    f.write_text("# Projet\nContenu mis à jour.")
    repo.index.add(["README.md"])
    repo.index.commit("fix: correction contenu")

    docs = GitExtractor(tmp_path).extract()
    commit_docs = [d for d in docs if d.metadonnees.get("sha")]

    assert len(commit_docs) >= 2
    for doc in commit_docs:
        assert_raw_document(doc, "git")
        assert doc.auteur == "Test User <test@example.com>"
        assert doc.metadonnees["sha"]
        assert len(doc.metadonnees["sha"]) == 40


def test_git_extract_branches_et_tags(tmp_path: Path) -> None:
    repo = _init_git_repo(tmp_path)

    f = tmp_path / "file.txt"
    f.write_text("contenu")
    repo.index.add(["file.txt"])
    repo.index.commit("feat: premier commit")

    repo.create_head("feature/test-extraction")
    repo.create_tag("v1.0.0", message="Release v1.0.0")

    docs = GitExtractor(tmp_path).extract()

    tag_docs = [d for d in docs if d.id_source.startswith("tag:")]
    assert len(tag_docs) >= 1
    assert any("v1.0.0" in d.titre for d in tag_docs)


def test_git_repo_invalide_retourne_liste_vide(tmp_path: Path) -> None:
    docs = GitExtractor(tmp_path / "inexistant").extract()
    assert docs == []


def test_git_git_non_trouve_leve_runtime_error(tmp_path: Path) -> None:
    with patch("kb_smart_metering.extractors.git.shutil.which", return_value=None):
        with patch.dict("os.environ", {}, clear=False):
            # S'assurer que GIT_PYTHON_GIT_EXECUTABLE n'est pas dans l'env
            import os
            os.environ.pop("GIT_PYTHON_GIT_EXECUTABLE", None)
            with pytest.raises(RuntimeError, match="git n'est pas trouvé"):
                GitExtractor(tmp_path)


# ---------------------------------------------------------------------------
# WordExtractor (office.py)
# ---------------------------------------------------------------------------


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    import docx

    doc = docx.Document()
    doc.add_heading("Titre principal", 0)
    doc.add_heading("Section 1", 1)
    doc.add_paragraph("Premier paragraphe de texte.")
    doc.add_paragraph("Deuxième paragraphe.")
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Col A"
    table.cell(0, 1).text = "Col B"
    table.cell(0, 2).text = "Col C"
    table.cell(1, 0).text = "val1"
    table.cell(1, 1).text = "val2"
    table.cell(1, 2).text = "val3"
    path = tmp_path / "document_test.docx"
    doc.save(str(path))
    return path


def test_word_extract_contenu(sample_docx: Path) -> None:
    docs = WordExtractor(sample_docx).extract()

    assert len(docs) == 1
    doc = docs[0]
    assert_raw_document(doc, "docx")
    assert "Titre principal" in doc.contenu_texte
    assert "Premier paragraphe" in doc.contenu_texte
    assert "Col A" in doc.contenu_texte
    assert doc.metadonnees["nb_tableaux"] == 1
    assert doc.metadonnees["nb_paragraphes"] >= 2


def test_word_titre_egal_nom_fichier(sample_docx: Path) -> None:
    docs = WordExtractor(sample_docx).extract()
    assert docs[0].titre == "document_test"


def test_word_fichier_inexistant_retourne_liste_vide(tmp_path: Path) -> None:
    assert WordExtractor(tmp_path / "inexistant.docx").extract() == []


# ---------------------------------------------------------------------------
# ExcelExtractor (office.py)
# ---------------------------------------------------------------------------


@pytest.fixture()
def sample_xlsx(tmp_path: Path) -> Path:
    import openpyxl

    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "Données"
    ws1.append(["Projet", "Statut", "Responsable"])
    ws1.append(["MDM v3", "En cours", "Jean Dupont"])
    ws1.append(["AMI 2.0", "Terminé", "Alice Martin"])
    ws2 = wb.create_sheet("Budget")
    ws2.append(["Poste", "Montant"])
    ws2.append(["Développement", "50000"])
    path = tmp_path / "classeur_test.xlsx"
    wb.save(str(path))
    return path


def test_excel_extract_une_entree_par_feuille(sample_xlsx: Path) -> None:
    docs = ExcelExtractor(sample_xlsx).extract()

    assert len(docs) == 2
    noms = {d.metadonnees["sheet_name"] for d in docs}
    assert "Données" in noms
    assert "Budget" in noms


def test_excel_contenu_en_markdown(sample_xlsx: Path) -> None:
    docs = ExcelExtractor(sample_xlsx).extract()
    donnees = next(d for d in docs if d.metadonnees["sheet_name"] == "Données")

    assert_raw_document(donnees, "xlsx")
    assert "Projet" in donnees.contenu_texte
    assert "MDM v3" in donnees.contenu_texte
    assert "|" in donnees.contenu_texte  # format Markdown


def test_excel_fichier_inexistant_retourne_liste_vide(tmp_path: Path) -> None:
    assert ExcelExtractor(tmp_path / "inexistant.xlsx").extract() == []


# ---------------------------------------------------------------------------
# PdfExtractor (office.py)
# ---------------------------------------------------------------------------


@pytest.fixture()
def sample_pdf(tmp_path: Path) -> Path:
    import fitz  # PyMuPDF

    pdf_doc = fitz.open()
    p1 = pdf_doc.new_page(width=595, height=842)
    p1.insert_text((72, 100), "Contenu de la première page du rapport.")
    p2 = pdf_doc.new_page(width=595, height=842)
    p2.insert_text((72, 100), "Contenu de la deuxième page avec des données.")
    path = tmp_path / "rapport_test.pdf"
    pdf_doc.save(str(path))
    pdf_doc.close()
    return path


def test_pdf_extract_une_entree_par_page(sample_pdf: Path) -> None:
    docs = PdfExtractor(sample_pdf).extract()

    assert len(docs) == 2
    pages = sorted(docs, key=lambda d: d.metadonnees["page"])
    assert_raw_document(pages[0], "pdf")
    assert pages[0].metadonnees["page"] == 1
    assert pages[1].metadonnees["page"] == 2
    assert pages[0].metadonnees["total_pages"] == 2


def test_pdf_contenu_avec_numero_page(sample_pdf: Path) -> None:
    docs = PdfExtractor(sample_pdf).extract()
    pages = sorted(docs, key=lambda d: d.metadonnees["page"])

    assert "[Page 1]" in pages[0].contenu_texte
    assert "première page" in pages[0].contenu_texte


def test_pdf_fichier_inexistant_retourne_liste_vide(tmp_path: Path) -> None:
    assert PdfExtractor(tmp_path / "inexistant.pdf").extract() == []


# ---------------------------------------------------------------------------
# MeetingExtractor
# ---------------------------------------------------------------------------


def test_meeting_extract_txt_fixture() -> None:
    """Teste le parsing depuis le fichier de fixture TXT."""
    extractor = MeetingExtractor(FIXTURES_DIR / "reunion_2024-01-15.txt")
    docs = extractor.extract()

    assert len(docs) == 1
    doc = docs[0]
    assert_raw_document(doc, "meeting")
    assert "Jean Dupont" in doc.contenu_texte
    assert "Alice Martin" in doc.contenu_texte
    participants = doc.metadonnees["participants"]
    assert "Jean Dupont" in participants
    assert "Alice Martin" in participants
    assert "Bob Durand" in participants
    assert doc.metadonnees["nb_segments"] >= 5
    assert doc.metadonnees["format"] == "txt"


def test_meeting_extract_vtt_fixture() -> None:
    """Teste le parsing depuis le fichier de fixture VTT."""
    extractor = MeetingExtractor(FIXTURES_DIR / "reunion_teams.vtt")
    docs = extractor.extract()

    assert len(docs) == 1
    doc = docs[0]
    assert_raw_document(doc, "meeting")
    assert "Jean Dupont" in doc.contenu_texte
    assert "Alice Martin" in doc.contenu_texte
    participants = doc.metadonnees["participants"]
    assert "Jean Dupont" in participants
    assert "Alice Martin" in participants
    assert doc.metadonnees["format"] == "vtt"
    assert doc.metadonnees["nb_segments"] >= 3


def test_meeting_timestamps_presents_dans_contenu(tmp_path: Path) -> None:
    fichier = tmp_path / "reunion.txt"
    fichier.write_text(
        "[00:01:30] Alice Martin: Premier point à l'ordre du jour.\n"
        "[00:02:00] Bob Durand: Merci Alice.\n",
        encoding="utf-8",
    )
    docs = MeetingExtractor(fichier).extract()
    assert len(docs) == 1
    assert "00:01:30" in docs[0].contenu_texte
    assert "Alice Martin" in docs[0].contenu_texte


def test_meeting_fichier_inexistant_retourne_liste_vide(tmp_path: Path) -> None:
    assert MeetingExtractor(tmp_path / "inexistant.txt").extract() == []


# ---------------------------------------------------------------------------
# Rétrocompatibilité files.py → office.py
# ---------------------------------------------------------------------------


def test_files_module_retro_compat() -> None:
    """Les imports depuis files.py continuent de fonctionner."""
    from kb_smart_metering.extractors.files import (  # noqa: F401
        ExcelExtractor as E,
        PdfExtractor as P,
        WordExtractor as W,
    )

    assert W is WordExtractor
    assert E is ExcelExtractor
    assert P is PdfExtractor

